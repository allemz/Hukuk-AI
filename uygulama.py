# -*- coding: utf-8 -*-
"""
Hukuk AI Asistanı - RAG Tabanlı Yapay Zeka Hukuk Sistemi
Groq API ile Çalışan Kapsamlı Hukuk Asistanı
"""

import streamlit as st
import os
import json
import hashlib
from datetime import datetime
from pathlib import Path
import chromadb
from chromadb.config import Settings
from groq import Groq
import PyPDF2
from docx import Document
import io
import base64
import asyncio
import httpx
from typing import Optional, List, Dict, Any
from pydantic import BaseModel, Field, HttpUrl, ConfigDict
import html
import re
from markitdown import MarkItDown
import logging

# Logging ayarları
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Sayfa yapılandırması
st.set_page_config(
    page_title="Hukuk AI Asistanı",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Stilleri
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #1e3a8a;
        margin-bottom: 1rem;
    }
    .info-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #f0f9ff;
        border-left: 4px solid #3b82f6;
        margin: 1rem 0;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #f0fdf4;
        border-left: 4px solid #22c55e;
        margin: 1rem 0;
    }
    .folder-card {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #ffffff;
        border: 1px solid #e5e7eb;
        margin: 0.5rem 0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    .doc-card {
        padding: 0.75rem;
        border-radius: 0.5rem;
        background-color: #fafafa;
        border: 1px solid #e5e7eb;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# EMSAL API - Models
# ============================================================================

class EmsalDetailedSearchRequestData(BaseModel):
    """Emsal detaylı arama için API payload modeli."""
    arananKelime: Optional[str] = ""
    Bam_Hukuk_Mahkemeleri: str = Field("", alias="Bam Hukuk Mahkemeleri")
    Hukuk_Mahkemeleri: str = Field("", alias="Hukuk Mahkemeleri")
    birimHukukMah: Optional[str] = Field("", description="Bölge adliye mahkemeleri daireleri")
    esasYil: Optional[str] = ""
    esasIlkSiraNo: Optional[str] = ""
    esasSonSiraNo: Optional[str] = ""
    kararYil: Optional[str] = ""
    kararIlkSiraNo: Optional[str] = ""
    kararSonSiraNo: Optional[str] = ""
    baslangicTarihi: Optional[str] = ""
    bitisTarihi: Optional[str] = ""
    siralama: str
    siralamaDirection: str
    pageSize: int
    pageNumber: int
    model_config = ConfigDict(populate_by_name=True)

class EmsalSearchRequest(BaseModel):
    """Emsal arama isteği modeli."""
    keyword: str = Field("", description="Arama kelimesi")
    selected_bam_civil_court: str = Field("", description="BAM Hukuk Mahkemesi")
    selected_civil_court: str = Field("", description="Hukuk Mahkemesi")
    selected_regional_civil_chambers: List[str] = Field(default_factory=list, description="Bölge daireleri")
    case_year_esas: str = Field("", description="Esas yılı")
    case_start_seq_esas: str = Field("", description="Esas başlangıç no")
    case_end_seq_esas: str = Field("", description="Esas bitiş no")
    decision_year_karar: str = Field("", description="Karar yılı")
    decision_start_seq_karar: str = Field("", description="Karar başlangıç no")
    decision_end_seq_karar: str = Field("", description="Karar bitiş no")
    start_date: str = Field("", description="Başlangıç tarihi (DD.MM.YYYY)")
    end_date: str = Field("", description="Bitiş tarihi (DD.MM.YYYY)")
    sort_criteria: str = Field("1", description="Sıralama")
    sort_direction: str = Field("desc", description="Yön")
    page_number: int = Field(default=1, ge=1)
    page_size: int = Field(default=10, ge=1, le=10)

class EmsalApiDecisionEntry(BaseModel):
    """Emsal API'den gelen karar girdisi."""
    id: str
    daire: str = Field("", description="Daire")
    esasNo: str = Field("", description="Esas No")
    kararNo: str = Field("", description="Karar No")
    kararTarihi: str = Field("", description="Karar Tarihi")
    arananKelime: str = Field("", description="Aranan Kelime")
    durum: str = Field("", description="Durum")
    document_url: Optional[HttpUrl] = Field(None, description="Belge URL")
    model_config = ConfigDict(extra='ignore')

class EmsalApiResponseInnerData(BaseModel):
    """Emsal API yanıt verisi."""
    data: List[EmsalApiDecisionEntry]
    recordsTotal: int
    recordsFiltered: int
    draw: int = Field(0, description="Çizim sayacı")

class EmsalApiResponse(BaseModel):
    """Emsal API tam yanıt modeli."""
    data: EmsalApiResponseInnerData
    metadata: Optional[Dict[str, Any]] = Field(None, description="Metadata")

class EmsalDocumentMarkdown(BaseModel):
    """Emsal karar belgesi (Markdown formatında)."""
    id: str
    markdown_content: str = Field("", description="Markdown içerik")
    source_url: HttpUrl

class CompactEmsalSearchResult(BaseModel):
    """Kompakt Emsal arama sonucu."""
    decisions: List[EmsalApiDecisionEntry]
    total_records: int
    requested_page: int
    page_size: int

# ============================================================================
# EMSAL API - Client
# ============================================================================

class EmsalApiClient:
    """UYAP Emsal (İçtihat) arama sistemi için API client."""

    BASE_URL = "https://emsal.uyap.gov.tr"
    DETAILED_SEARCH_ENDPOINT = "/aramadetaylist"
    DOCUMENT_ENDPOINT = "/getDokuman"

    def __init__(self, request_timeout: float = 30.0):
        self.http_client = httpx.AsyncClient(
            base_url=self.BASE_URL,
            headers={
                "Content-Type": "application/json; charset=UTF-8",
                "Accept": "application/json, text/plain, */*",
                "X-Requested-With": "XMLHttpRequest",
            },
            timeout=request_timeout,
            verify=False
        )

    async def search_detailed_decisions(self, params: EmsalSearchRequest) -> EmsalApiResponse:
        """Emsal sisteminde detaylı arama yapar."""
        data_for_api_payload = EmsalDetailedSearchRequestData(
            arananKelime=params.keyword or "",
            Bam_Hukuk_Mahkemeleri=params.selected_bam_civil_court,
            Hukuk_Mahkemeleri=params.selected_civil_court,
            birimHukukMah="+".join(params.selected_regional_civil_chambers) if params.selected_regional_civil_chambers else "",
            esasYil=params.case_year_esas or "",
            esasIlkSiraNo=params.case_start_seq_esas or "",
            esasSonSiraNo=params.case_end_seq_esas or "",
            kararYil=params.decision_year_karar or "",
            kararIlkSiraNo=params.decision_start_seq_karar or "",
            kararSonSiraNo=params.decision_end_seq_karar or "",
            baslangicTarihi=params.start_date or "",
            bitisTarihi=params.end_date or "",
            siralama=params.sort_criteria,
            siralamaDirection=params.sort_direction,
            pageSize=params.page_size,
            pageNumber=params.page_number
        )

        payload_dict = data_for_api_payload.model_dump(by_alias=True, exclude_none=True)
        cleaned_payload = {k: v for k, v in payload_dict.items() if v != ""}
        final_payload = {"data": cleaned_payload}

        logger.info(f"Emsal araması yapılıyor: {params.keyword}")
        return await self._execute_api_search(self.DETAILED_SEARCH_ENDPOINT, final_payload)

    async def _execute_api_search(self, endpoint: str, payload: Dict) -> EmsalApiResponse:
        """API arama isteği yapar ve yanıtı işler."""
        try:
            response = await self.http_client.post(endpoint, json=payload)
            response.raise_for_status()
            response_json_data = response.json()

            api_response_parsed = EmsalApiResponse(**response_json_data)

            if api_response_parsed.data and api_response_parsed.data.data:
                for decision_item in api_response_parsed.data.data:
                    if decision_item.id:
                        decision_item.document_url = f"{self.BASE_URL}{self.DOCUMENT_ENDPOINT}?id={decision_item.id}"

            return api_response_parsed
        except httpx.RequestError as e:
            logger.error(f"Emsal API bağlantı hatası: {e}")
            raise
        except Exception as e:
            logger.error(f"Emsal API yanıt hatası: {e}")
            raise

    def _clean_html_and_convert_to_markdown(self, html_content: str) -> Optional[str]:
        """HTML içeriği temizler ve Markdown'a çevirir."""
        if not html_content:
            return None

        content = html.unescape(html_content)
        content = content.replace('\"', '"')
        content = content.replace('\r\n', '\n')
        content = content.replace('\n', '\n')
        content = content.replace('\t', '\t')

        markdown_text = None
        try:
            html_bytes = content.encode('utf-8')
            html_stream = io.BytesIO(html_bytes)

            md_converter = MarkItDown()
            conversion_result = md_converter.convert(html_stream)
            markdown_text = conversion_result.text_content
            logger.info("HTML -> Markdown dönüşümü başarılı")
        except Exception as e:
            logger.error(f"Markdown dönüşüm hatası: {e}")

        return markdown_text

    async def get_decision_document_as_markdown(self, id: str) -> EmsalDocumentMarkdown:
        """Belirli bir Emsal kararını ID ile alır ve Markdown formatında döndürür."""
        document_api_url = f"{self.DOCUMENT_ENDPOINT}?id={id}"
        source_url = f"{self.BASE_URL}{document_api_url}"
        logger.info(f"Emsal belgesi alınıyor (ID: {id})")

        try:
            response = await self.http_client.get(document_api_url)
            response.raise_for_status()

            response_json = response.json()
            html_content = response_json.get("data")

            if not isinstance(html_content, str) or not html_content.strip():
                logger.warning(f"Emsal belgesi boş (ID: {id})")
                return EmsalDocumentMarkdown(id=id, markdown_content="", source_url=source_url)

            markdown_content = self._clean_html_and_convert_to_markdown(html_content)

            return EmsalDocumentMarkdown(
                id=id,
                markdown_content=markdown_content or "",
                source_url=source_url
            )
        except Exception as e:
            logger.error(f"Emsal belge alma hatası (ID: {id}): {e}")
            raise

    async def close_client_session(self):
        """HTTP client oturumunu kapatır."""
        if self.http_client and not self.http_client.is_closed:
            await self.http_client.aclose()
        logger.info("Emsal API client kapatıldı")

# ============================================================================
# HUKUK AI ASİSTANI - Ana Sınıf
# ============================================================================

class HukukAIAsistani:
    """Hukuk AI Asistanı Ana Sınıfı"""

    def __init__(self):
        """Sistem başlatma"""
        self.groq_api_key = st.session_state.get('groq_api_key', '')
        self.data_dir = Path("hukuk_data")
        self.data_dir.mkdir(exist_ok=True)

        # ChromaDB başlatma
        self.chroma_client = chromadb.PersistentClient(
            path=str(self.data_dir / "chroma_db"),
            settings=Settings(anonymized_telemetry=False)
        )

        # Koleksiyonlar
        self.collections = {
            'ictihatlar': self._get_or_create_collection('ictihatlar'),
            'mevzuat': self._get_or_create_collection('mevzuat'),
            'sozlesmeler': self._get_or_create_collection('sozlesmeler'),
            'dilekce': self._get_or_create_collection('dilekce'),
        }

        # Emsal API Client
        self.emsal_client = EmsalApiClient(request_timeout=30.0)

    def _get_or_create_collection(self, name):
        """Koleksiyon oluştur veya getir"""
        try:
            return self.chroma_client.get_or_create_collection(
                name=name,
                metadata={"hnsw:space": "cosine"}
            )
        except Exception as e:
            st.error(f"Koleksiyon hatası: {e}")
            return None

    def get_groq_client(self):
        """Groq API istemcisi oluştur"""
        if not self.groq_api_key:
            raise ValueError("Groq API anahtarı gerekli")
        return Groq(api_key=self.groq_api_key)

    def generate_embedding(self, text):
        """Metin için embedding oluştur (basit hash tabanlı)"""
        hash_obj = hashlib.sha256(text.encode())
        hash_hex = hash_obj.hexdigest()
        embedding = [int(hash_hex[i:i+2], 16) / 255.0 for i in range(0, min(len(hash_hex), 768), 2)]
        while len(embedding) < 384:
            embedding.append(0.0)
        return embedding[:384]

    def add_document(self, collection_name, document_text, metadata):
        """Belge ekle"""
        try:
            collection = self.collections.get(collection_name)
            if not collection:
                return False

            doc_id = hashlib.md5(
                (document_text + str(datetime.now())).encode()
            ).hexdigest()

            chunks = [document_text[i:i+1000] for i in range(0, len(document_text), 1000)]

            for idx, chunk in enumerate(chunks):
                chunk_id = f"{doc_id}_chunk_{idx}"
                embedding = self.generate_embedding(chunk)

                collection.add(
                    ids=[chunk_id],
                    documents=[chunk],
                    embeddings=[embedding],
                    metadatas=[{**metadata, 'chunk_index': idx, 'parent_id': doc_id}]
                )

            return True
        except Exception as e:
            st.error(f"Belge ekleme hatası: {e}")
            return False

    def search_documents(self, collection_name, query, n_results=5):
        """Belge ara"""
        try:
            collection = self.collections.get(collection_name)
            if not collection:
                return []

            query_embedding = self.generate_embedding(query)

            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results
            )

            return results
        except Exception as e:
            st.error(f"Arama hatası: {e}")
            return {'documents': [[]], 'metadatas': [[]], 'distances': [[]]}

    def chat_with_rag(self, user_query, collection_names=['ictihatlar', 'mevzuat']):
        """RAG tabanlı sohbet"""
        try:
            client = self.get_groq_client()

            context_parts = []
            for col_name in collection_names:
                results = self.search_documents(col_name, user_query, n_results=3)
                if results and results.get('documents') and results['documents'][0]:
                    for doc in results['documents'][0]:
                        context_parts.append(doc)

            context = "\n\n---\n\n".join(context_parts) if context_parts else "İlgili belge bulunamadı."

            system_prompt = """Sen uzman bir hukuk asistanısın. Türkiye hukuk sistemine hakimsin.
Verilen bağlam bilgilerini kullanarak detaylı ve profesyonel hukuki cevaplar ver.
Cevaplarını madde madde ve net bir şekilde yaz. Kaynak göstermeyi unutma."""

            user_prompt = f"""Bağlam Bilgileri:
{context}

---

Kullanıcı Sorusu: {user_query}

Lütfen yukarıdaki bağlam bilgilerini dikkate alarak detaylı bir hukuki değerlendirme yap."""

            completion = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=2000,
                top_p=0.9,
                stream=False
            )

            return completion.choices[0].message.content

        except Exception as e:
            return f"Hata oluştu: {str(e)}"

    def generate_document(self, doc_type, prompt):
        """Belge oluştur"""
        try:
            client = self.get_groq_client()

            templates = {
                'sozlesme': "Sen profesyonel bir hukuk uzmanısın. Aşağıdaki talebe göre detaylı bir sözleşme metni hazırla. Sözleşme Türkiye hukuk sistemine uygun olmalı ve tüm gerekli maddeleri içermelidir.",
                'dilekce': "Sen deneyimli bir avukatsın. Aşağıdaki talebe göre mahkemeye sunulacak profesyonel bir dilekçe hazırla. Dilekçe tüm hukuki gereklilikleri karşılamalı ve düzgün formatlanmış olmalıdır.",
                'dava': "Sen hukuk profesyonelinden dava hazırlama konusunda yardım al. Detaylı ve kapsamlı bir dava dosyası hazırla."
            }

            system_prompt = templates.get(doc_type, templates['sozlesme'])

            completion = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,
                max_tokens=3000,
                top_p=0.9,
                stream=False
            )

            return completion.choices[0].message.content

        except Exception as e:
            return f"Belge oluşturma hatası: {str(e)}"

    def analyze_document(self, document_text, analysis_type):
        """Belge analizi"""
        try:
            client = self.get_groq_client()

            analysis_prompts = {
                'genel': "Bu belgeyi detaylı analiz et. Güçlü ve zayıf yönlerini, eksiklikleri ve önerilerini belirt.",
                'risk': "Bu belgedeki hukuki riskleri ve potansiyel sorunları analiz et.",
                'madde': "Bu belgedeki maddeleri tek tek incele ve değerlendir."
            }

            prompt = f"""{analysis_prompts.get(analysis_type, analysis_prompts['genel'])}

Belge:
{document_text}"""

            completion = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": "Sen uzman bir hukuk analisti ve danışmanısın."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2500,
                top_p=0.9,
                stream=False
            )

            return completion.choices[0].message.content

        except Exception as e:
            return f"Analiz hatası: {str(e)}"

    async def search_emsal_decisions(self, keyword: str, year: str = "", page: int = 1, page_size: int = 10):
        """UYAP Emsal sisteminde karar arama yapar."""
        try:
            search_params = EmsalSearchRequest(
                keyword=keyword,
                selected_bam_civil_court="",
                selected_civil_court="",
                selected_regional_civil_chambers=[],
                decision_year_karar=year,
                sort_criteria="1",
                sort_direction="desc",
                page_number=page,
                page_size=page_size
            )

            response = await self.emsal_client.search_detailed_decisions(search_params)

            return {
                'success': True,
                'decisions': response.data.data,
                'total_records': response.data.recordsTotal,
                'filtered_records': response.data.recordsFiltered
            }
        except Exception as e:
            logger.error(f"Emsal arama hatası: {e}")
            return {
                'success': False,
                'error': str(e),
                'decisions': [],
                'total_records': 0
            }

    async def get_emsal_decision_content(self, decision_id: str):
        """Emsal karar içeriğini alır."""
        try:
            document = await self.emsal_client.get_decision_document_as_markdown(decision_id)
            return {
                'success': True,
                'content': document.markdown_content,
                'source_url': str(document.source_url)
            }
        except Exception as e:
            logger.error(f"Emsal belge alma hatası: {e}")
            return {
                'success': False,
                'error': str(e),
                'content': ''
            }

    async def add_emsal_decision_to_db(self, decision_id: str, metadata: dict):
        """Emsal kararını veritabanına ekler."""
        try:
            doc_result = await self.get_emsal_decision_content(decision_id)

            if doc_result['success'] and doc_result['content']:
                success = self.add_document(
                    'ictihatlar',
                    doc_result['content'],
                    {
                        **metadata,
                        'source': 'UYAP Emsal',
                        'source_url': doc_result['source_url']
                    }
                )
                return success
            return False
        except Exception as e:
            logger.error(f"Emsal veritabanına ekleme hatası: {e}")
            return False

def extract_text_from_pdf(pdf_file):
    """PDF'den metin çıkar"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"PDF okuma hatası: {str(e)}"

def extract_text_from_docx(docx_file):
    """DOCX'ten metin çıkar"""
    try:
        doc = Document(docx_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"DOCX okuma hatası: {str(e)}"

def save_folders():
    """Klasörleri kaydet"""
    if 'folders' not in st.session_state:
        st.session_state.folders = {}

    folder_file = Path("hukuk_data/folders.json")
    with open(folder_file, 'w', encoding='utf-8') as f:
        json.dump(st.session_state.folders, f, ensure_ascii=False, indent=2)

def load_folders():
    """Klasörleri yükle"""
    folder_file = Path("hukuk_data/folders.json")
    if folder_file.exists():
        with open(folder_file, 'r', encoding='utf-8') as f:
            st.session_state.folders = json.load(f)
    else:
        st.session_state.folders = {}

def main():
    """Ana uygulama"""

    # Başlık
    st.markdown('<h1 class="main-header">⚖️ Hukuk AI Asistanı</h1>', unsafe_allow_html=True)
    st.markdown("### Türkiye'nin En Kapsamlı AI Destekli Hukuk Platformu")

    # Sidebar - API Ayarları
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/law.png", width=80)
        st.title("⚙️ Ayarlar")

        api_key = st.text_input(
            "Groq API Anahtarı",
            type="password",
            value=st.session_state.get('groq_api_key', ''),
            help="Groq API anahtarınızı girin (https://console.groq.com/)"
        )

        if api_key:
            st.session_state.groq_api_key = api_key
            st.success("✅ API anahtarı kaydedildi")

        st.divider()

        st.markdown("### 📊 Sistem İstatistikleri")
        if 'ai_assistant' in st.session_state:
            try:
                ictihat_count = st.session_state.ai_assistant.collections['ictihatlar'].count()
                mevzuat_count = st.session_state.ai_assistant.collections['mevzuat'].count()
                st.metric("İçtihat", ictihat_count)
                st.metric("Mevzuat", mevzuat_count)
            except:
                st.info("Veritabanı yükleniyor...")

    # Ana sistem başlatma
    if 'ai_assistant' not in st.session_state and st.session_state.get('groq_api_key'):
        st.session_state.ai_assistant = HukukAIAsistani()
        load_folders()

    # API anahtarı kontrolü
    if not st.session_state.get('groq_api_key'):
        st.warning("⚠️ Lütfen sol menüden Groq API anahtarınızı girin.")
        st.info("""
        **API Anahtarı Nasıl Alınır?**
        1. https://console.groq.com/ adresine gidin
        2. Hesap oluşturun veya giriş yapın
        3. API Keys bölümünden yeni anahtar oluşturun
        4. Anahtarı kopyalayıp sol menüye yapıştırın
        """)
        return

    # Ana Tabs
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "🤖 AI Asistan",
        "🔍 Karar Arama",
        "⚖️ UYAP Emsal",
        "📝 Belge Yazım",
        "📂 Dava Yönetimi",
        "📚 Veritabanı"
    ])

    # Tab 1: AI Asistan
    with tab1:
        st.header("🤖 Yapay Zeka Hukuk Asistanı")
        st.markdown("""
        <div class="info-box">
        ✨ Özellikler:<br>
        • Her türlü hukuki sorunuza anında cevap<br>
        • Sözleşme ve dilekçe analizi<br>
        • Hukuki araştırma desteği<br>
        • İçtihat ve mevzuat bilgisiyle desteklenen yanıtlar
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([2, 1])

        with col1:
            user_query = st.text_area(
                "Hukuki sorunuzu sorun:",
                height=100,
                placeholder="Örn: İş sözleşmesinde dikkat edilmesi gereken maddeler nelerdir?"
            )

            col_a, col_b = st.columns(2)
            with col_a:
                search_ictihat = st.checkbox("İçtihatlar", value=True)
            with col_b:
                search_mevzuat = st.checkbox("Mevzuat", value=True)

            if st.button("🚀 Sor", type="primary", use_container_width=True):
                if user_query:
                    with st.spinner("AI düşünüyor..."):
                        collections = []
                        if search_ictihat:
                            collections.append('ictihatlar')
                        if search_mevzuat:
                            collections.append('mevzuat')

                        response = st.session_state.ai_assistant.chat_with_rag(
                            user_query,
                            collection_names=collections
                        )

                        st.markdown("### 📋 Yanıt:")
                        st.markdown(response)
                else:
                    st.warning("Lütfen bir soru girin.")

        with col2:
            st.markdown("### 💡 Örnek Sorular")
            example_questions = [
                "Kira sözleşmesinde nelere dikkat etmeliyim?",
                "İş sözleşmesi feshi şartları nelerdir?",
                "Boşanma davası nasıl açılır?",
                "Tüketici hakları nelerdir?",
                "Miras hukuku temel kavramları"
            ]

            for q in example_questions:
                if st.button(q, key=f"example_{q[:20]}"):
                    st.session_state.example_query = q
                    st.rerun()

    # Tab 2: Karar Arama  
    with tab2:
        st.header("🔍 Karar ve Mevzuat Arama Motoru")

        col1, col2 = st.columns([3, 1])

        with col1:
            search_query = st.text_input(
                "Semantik arama yapın:",
                placeholder="Örn: iş güvencesi tazminatı"
            )

        with col2:
            search_type = st.selectbox(
                "Arama Tipi:",
                ["İçtihat", "Mevzuat", "Tümü"]
            )

        if st.button("🔎 Ara", use_container_width=True):
            if search_query:
                with st.spinner("Aranıyor..."):
                    collections_to_search = []

                    if search_type in ["İçtihat", "Tümü"]:
                        collections_to_search.append('ictihatlar')
                    if search_type in ["Mevzuat", "Tümü"]:
                        collections_to_search.append('mevzuat')

                    all_results = []
                    for col_name in collections_to_search:
                        results = st.session_state.ai_assistant.search_documents(
                            col_name,
                            search_query,
                            n_results=5
                        )

                        if results and results.get('documents') and results['documents'][0]:
                            for idx, doc in enumerate(results['documents'][0]):
                                metadata = results['metadatas'][0][idx] if results.get('metadatas') else {}
                                distance = results['distances'][0][idx] if results.get('distances') else 0

                                all_results.append({
                                    'type': col_name,
                                    'text': doc,
                                    'metadata': metadata,
                                    'relevance': 1 - distance
                                })

                    if all_results:
                        st.success(f"✅ {len(all_results)} sonuç bulundu")

                        for idx, result in enumerate(all_results, 1):
                            with st.expander(f"📄 Sonuç {idx} - {result['type'].title()} (İlgililik: {result['relevance']:.2%})"):
                                st.markdown(f"**Tip:** {result['type'].title()}")
                                st.markdown(f"**Metadata:** {result['metadata']}")
                                st.markdown("**İçerik:**")
                                st.text(result['text'][:500] + "..." if len(result['text']) > 500 else result['text'])
                    else:
                        st.info("Sonuç bulunamadı. Lütfen farklı anahtar kelimeler deneyin.")
            else:
                st.warning("Lütfen arama terimi girin.")

    # Tab 3: UYAP Emsal
    with tab3:
        st.header("⚖️ UYAP Emsal - Canlı İçtihat Arama")
        st.markdown("""
        <div class="info-box">
        ✨ UYAP Emsal Özellikleri:<br>
        • Türkiye'nin resmi içtihat sistemine doğrudan erişim<br>
        • Yargıtay ve Bölge Adliye Mahkemeleri kararları<br>
        • Anlık karar arama ve içerik görüntüleme<br>
        • Kararları veritabanınıza otomatik ekleme
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([3, 1])

        with col1:
            emsal_keyword = st.text_input(
                "🔍 Arama Kelimesi:",
                placeholder="Örn: iş güvencesi, tazminat, kira artışı",
                key="emsal_keyword"
            )

        with col2:
            emsal_year = st.text_input(
                "📅 Karar Yılı:",
                placeholder="2023",
                key="emsal_year"
            )

        if st.button("🚀 UYAP Emsal'de Ara", type="primary", use_container_width=True):
            if emsal_keyword:
                with st.spinner("UYAP Emsal sisteminde aranıyor..."):
                    try:
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)

                        result = loop.run_until_complete(
                            st.session_state.ai_assistant.search_emsal_decisions(
                                keyword=emsal_keyword,
                                year=emsal_year,
                                page=1,
                                page_size=10
                            )
                        )

                        loop.close()

                        if result['success']:
                            st.success(f"✅ {result['total_records']} sonuç bulundu")

                            for idx, decision in enumerate(result['decisions'], 1):
                                with st.expander(f"📄 {decision.daire} - {decision.esasNo}"):
                                    st.markdown(f"**Karar Tarihi:** {decision.kararTarihi}")
                                    st.markdown(f"**Karar No:** {decision.kararNo}")
                        else:
                            st.error(f"❌ Arama başarısız: {result.get('error')}")
                    except Exception as e:
                        st.error(f"❌ Hata: {str(e)}")
            else:
                st.warning("⚠️ Lütfen arama kelimesi girin")

    # Tab 4: Belge Yazım
    with tab4:
        st.header("📝 Belge Yazım Asistanı ve Editörü")

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("✏️ Yeni Belge Oluştur")

            doc_type = st.selectbox(
                "Belge Tipi:",
                ["Sözleşme", "Dilekçe", "Dava Dosyası"]
            )

            doc_prompt = st.text_area(
                "Belge detaylarını açıklayın:",
                height=150,
                placeholder="Örn: İki şirket arasında gizlilik sözleşmesi hazırlamak istiyorum..."
            )

            if st.button("📄 Belge Oluştur", type="primary", use_container_width=True):
                if doc_prompt:
                    with st.spinner("Belge hazırlanıyor..."):
                        doc_type_map = {
                            'Sözleşme': 'sozlesme',
                            'Dilekçe': 'dilekce',
                            'Dava Dosyası': 'dava'
                        }

                        generated_doc = st.session_state.ai_assistant.generate_document(
                            doc_type_map[doc_type],
                            doc_prompt
                        )

                        st.session_state.generated_document = generated_doc
                        st.success("✅ Belge oluşturuldu!")
                else:
                    st.warning("Lütfen belge detaylarını girin.")

        with col2:
            st.subheader("🔍 Belge Analizi")

            uploaded_file = st.file_uploader(
                "Analiz edilecek belgeyi yükleyin:",
                type=['pdf', 'docx', 'txt']
            )

            if uploaded_file:
                analysis_type = st.selectbox(
                    "Analiz Tipi:",
                    ["Genel Analiz", "Risk Analizi", "Madde İncelemesi"]
                )

                if st.button("🔬 Analiz Et", use_container_width=True):
                    with st.spinner("Belge analiz ediliyor..."):
                        if uploaded_file.name.endswith('.pdf'):
                            doc_text = extract_text_from_pdf(uploaded_file)
                        elif uploaded_file.name.endswith('.docx'):
                            doc_text = extract_text_from_docx(uploaded_file)
                        else:
                            doc_text = uploaded_file.read().decode('utf-8')

                        analysis_type_map = {
                            'Genel Analiz': 'genel',
                            'Risk Analizi': 'risk',
                            'Madde İncelemesi': 'madde'
                        }

                        analysis = st.session_state.ai_assistant.analyze_document(
                            doc_text,
                            analysis_type_map[analysis_type]
                        )

                        st.session_state.document_analysis = analysis
                        st.success("✅ Analiz tamamlandı!")

        # Oluşturulan belgeyi göster
        if 'generated_document' in st.session_state:
            st.divider()
            st.subheader("📋 Oluşturulan Belge")

            edited_doc = st.text_area(
                "Belgeyi düzenleyin:",
                value=st.session_state.generated_document,
                height=400
            )

            col_a, col_b, col_c = st.columns(3)

            with col_a:
                if st.button("💾 DOCX Olarak İndir"):
                    doc = Document()
                    for paragraph in edited_doc.split('\n'):
                        doc.add_paragraph(paragraph)

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        "📥 İndir",
                        buffer,
                        file_name=f"belge_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            with col_b:
                if st.button("📄 PDF Olarak İndir"):
                    st.info("PDF oluşturma özelliği yakında eklenecek")

            with col_c:
                if st.button("📂 Klasöre Kaydet"):
                    st.session_state.save_to_folder = edited_doc

        # Analiz sonucunu göster
        if 'document_analysis' in st.session_state:
            st.divider()
            st.subheader("🔬 Analiz Sonucu")
            st.markdown(st.session_state.document_analysis)

    # Tab 5: Dava Yönetimi
    with tab5:
        st.header("📂 Klasörler ile Dava Yönetimi")

        col1, col2 = st.columns([1, 2])

        with col1:
            st.subheader("📁 Klasörler")

            # Yeni klasör oluştur
            with st.expander("➕ Yeni Klasör"):
                new_folder_name = st.text_input("Klasör Adı:", key="new_folder")
                folder_desc = st.text_area("Açıklama:", key="folder_desc")

                if st.button("Oluştur"):
                    if new_folder_name:
                        if 'folders' not in st.session_state:
                            st.session_state.folders = {}

                        folder_id = hashlib.md5(new_folder_name.encode()).hexdigest()[:8]
                        st.session_state.folders[folder_id] = {
                            'name': new_folder_name,
                            'description': folder_desc,
                            'created': datetime.now().isoformat(),
                            'documents': [],
                            'notes': []
                        }
                        save_folders()
                        st.success(f"✅ Klasör '{new_folder_name}' oluşturuldu")
                        st.rerun()

            # Klasör listesi
            if 'folders' in st.session_state and st.session_state.folders:
                for folder_id, folder_data in st.session_state.folders.items():
                    with st.container():
                        st.markdown(f"""
                        <div class="folder-card">
                        📁 <strong>{folder_data['name']}</strong><br>
                        <small>{len(folder_data.get('documents', []))} belge</small>
                        </div>
                        """, unsafe_allow_html=True)

                        if st.button(f"Aç", key=f"open_{folder_id}"):
                            st.session_state.active_folder = folder_id
                            st.rerun()
            else:
                st.info("Henüz klasör yok. Yeni klasör oluşturun.")

        with col2:
            if 'active_folder' in st.session_state:
                folder_id = st.session_state.active_folder
                folder = st.session_state.folders.get(folder_id, {})

                st.subheader(f"📂 {folder.get('name', 'Klasör')}")
                st.caption(folder.get('description', ''))

                # Belge yükleme
                with st.expander("📤 Belge Yükle"):
                    uploaded_doc = st.file_uploader(
                        "Dosya seçin:",
                        type=['pdf', 'docx', 'txt', 'jpg', 'png'],
                        key=f"upload_{folder_id}"
                    )

                    doc_note = st.text_input("Not:", key=f"note_{folder_id}")

                    if st.button("Ekle", key=f"add_{folder_id}"):
                        if uploaded_doc:
                            doc_data = {
                                'filename': uploaded_doc.name,
                                'type': uploaded_doc.type,
                                'added': datetime.now().isoformat(),
                                'note': doc_note,
                                'size': uploaded_doc.size
                            }

                            if 'documents' not in folder:
                                folder['documents'] = []

                            folder['documents'].append(doc_data)
                            save_folders()
                            st.success(f"✅ '{uploaded_doc.name}' eklendi")
                            st.rerun()

                # Belgeler listesi
                st.subheader("📑 Belgeler")

                documents = folder.get('documents', [])
                if documents:
                    for idx, doc in enumerate(documents):
                        with st.container():
                            st.markdown(f"""
                            <div class="doc-card">
                            📄 <strong>{doc['filename']}</strong><br>
                            <small>{doc.get('note', '')}</small><br>
                            <small>Eklenme: {doc['added'][:10]}</small>
                            </div>
                            """, unsafe_allow_html=True)

                            if st.button("🗑️ Sil", key=f"del_{folder_id}_{idx}"):
                                documents.pop(idx)
                                save_folders()
                                st.rerun()
                else:
                    st.info("Henüz belge yok")

                # Notlar
                st.subheader("📝 Notlar")

                new_note = st.text_area("Yeni not:", key=f"new_note_{folder_id}")
                if st.button("Not Ekle", key=f"add_note_{folder_id}"):
                    if new_note:
                        if 'notes' not in folder:
                            folder['notes'] = []

                        folder['notes'].append({
                            'text': new_note,
                            'created': datetime.now().isoformat()
                        })
                        save_folders()
                        st.success("✅ Not eklendi")
                        st.rerun()

                notes = folder.get('notes', [])
                if notes:
                    for idx, note in enumerate(notes):
                        st.text_area(
                            f"Not {idx+1} ({note['created'][:10]}):",
                            value=note['text'],
                            key=f"note_view_{folder_id}_{idx}",
                            disabled=True
                        )
            else:
                st.info("👈 Sol menüden bir klasör seçin")

    # Tab 6: Veritabanı
    with tab6:
        st.header("📚 Veritabanı Yönetimi")

        tab6_1, tab6_2 = st.tabs(["Belge Ekle", "Toplu Yükleme"])

        with tab6_1:
            st.subheader("➕ Tekil Belge Ekle")

            col1, col2 = st.columns(2)

            with col1:
                collection = st.selectbox(
                    "Koleksiyon:",
                    ["İçtihatlar", "Mevzuat", "Sözleşmeler", "Dilekçe"]
                )

                doc_title = st.text_input("Belge Başlığı:")
                doc_category = st.text_input("Kategori:")

            with col2:
                doc_date = st.date_input("Tarih:")
                doc_source = st.text_input("Kaynak:")

            doc_content = st.text_area(
                "Belge İçeriği:",
                height=200,
                placeholder="Belge metnini buraya yapıştırın..."
            )

            if st.button("💾 Veritabanına Ekle", type="primary"):
                if doc_content and doc_title:
                    collection_map = {
                        'İçtihatlar': 'ictihatlar',
                        'Mevzuat': 'mevzuat',
                        'Sözleşmeler': 'sozlesmeler',
                        'Dilekçe': 'dilekce'
                    }

                    metadata = {
                        'title': doc_title,
                        'category': doc_category,
                        'date': str(doc_date),
                        'source': doc_source,
                        'added': datetime.now().isoformat()
                    }

                    success = st.session_state.ai_assistant.add_document(
                        collection_map[collection],
                        doc_content,
                        metadata
                    )

                    if success:
                        st.success(f"✅ Belge '{doc_title}' başarıyla eklendi!")
                    else:
                        st.error("❌ Belge eklenirken hata oluştu")
                else:
                    st.warning("Lütfen başlık ve içerik girin")

        with tab6_2:
            st.subheader("📦 Toplu Belge Yükleme")

            st.info("""
            **Toplu yükleme formatı:**
            - JSON dosyası yükleyin
            - Her belge şu formatta olmalı:
            ```json
            {
                "collection": "ictihatlar",
                "title": "Belge Başlığı",
                "content": "Belge içeriği...",
                "metadata": {"key": "value"}
            }
            ```
            """)

            bulk_file = st.file_uploader(
                "JSON dosyası seçin:",
                type=['json']
            )

            if bulk_file and st.button("📥 Toplu Yükle"):
                try:
                    data = json.load(bulk_file)

                    if isinstance(data, list):
                        success_count = 0
                        progress_bar = st.progress(0)

                        for idx, item in enumerate(data):
                            collection = item.get('collection', 'ictihatlar')
                            title = item.get('title', f'Belge {idx+1}')
                            content = item.get('content', '')
                            metadata = item.get('metadata', {})
                            metadata['title'] = title

                            if content:
                                success = st.session_state.ai_assistant.add_document(
                                    collection,
                                    content,
                                    metadata
                                )
                                if success:
                                    success_count += 1

                            progress_bar.progress((idx + 1) / len(data))

                        st.success(f"✅ {success_count}/{len(data)} belge başarıyla eklendi!")
                    else:
                        st.error("JSON dosyası liste formatında olmalı")

                except Exception as e:
                    st.error(f"Hata: {str(e)}")

if __name__ == "__main__":
    main()
